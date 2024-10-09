import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from statsmodels.regression.linear_model import OLS
from statsmodels.tools.tools import add_constant
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler
from scipy import stats
from scipy.interpolate import interp1d
from scipy.stats import binom
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split
from sklearn.metrics import classification_report
from io import BytesIO
from docx import Document
from docx.shared import Inches

# Function definitions_________________________________________________________________________

# Function for custom mathematical transformation
def apply_custom_transformation(data, transformation):
    if transformation == "Log (log)":
        return np.log(data)
    elif transformation == "Natural Log (ln)":
        return np.log1p(data)
    elif transformation == "Absolute Value (abs)":
        return np.abs(data)
    elif transformation == "Phase":
        return np.angle(data)
    else:
        return data  # No transformation if None selected

# Function to check normality assumption of residuals
def check_normality_residuals(residuals):
    st.write("Checking Normality of Residuals:")
    fig, ax = plt.subplots(1, 2, figsize=(12, 6))
    
    # Histogram of residuals
    ax[0].hist(residuals, bins=30, color='skyblue')
    ax[0].set_title('Histogram of Residuals')
    
    # Q-Q plot
    stats.probplot(residuals, dist="norm", plot=ax[1])
    ax[1].set_title('Q-Q Plot')
    
    st.pyplot(fig)

# Function to perform full data analysis
def full_data_analysis(x, y, x_label, y_label, threshold=None):
    

    # Ask user if they want to apply boundaries to input/output
    use_bounds_input = st.sidebar.checkbox(f"Set boundaries for {x_label}?")
    if use_bounds_input:
        lower_bound_input = st.sidebar.number_input(f"Lower bound for {x_label}", value=float(np.min(x)))
        upper_bound_input = st.sidebar.number_input(f"Upper bound for {x_label}", value=float(np.max(x)))
        x = np.clip(x, lower_bound_input, upper_bound_input)
        st.write(f"Applied boundaries to {x_label}: Lower = {lower_bound_input}, Upper = {upper_bound_input}")

    use_bounds_output = st.sidebar.checkbox(f"Set boundaries for {y_label}?")
    if use_bounds_output:
        lower_bound_output = st.sidebar.number_input(f"Lower bound for {y_label}", value=float(np.min(y)))
        upper_bound_output = st.sidebar.number_input(f"Upper bound for {y_label}", value=float(np.max(y)))
        y = np.clip(y, lower_bound_output, upper_bound_output)
        st.write(f"Applied boundaries to {y_label}: Lower = {lower_bound_output}, Upper = {upper_bound_output}")

    # Ask user if they want to apply transformations to input/output
    transformation_input = st.sidebar.selectbox(f"Apply transformation to {x_label}", ["None", "Log (log)", "Natural Log (ln)", "Absolute Value (abs)", "Phase"])
    if transformation_input != "None":
        x_label=f"{transformation_input}[{x_label}]"
        x = apply_custom_transformation(x, transformation_input)
        st.write(f"Applied {transformation_input} transformation to {x_label}")

    transformation_output = st.sidebar.selectbox(f"Apply transformation to {y_label}", ["None", "Log (log)", "Natural Log (ln)", "Absolute Value (abs)", "Phase"])
    if transformation_output != "None":
        y_label=f"{transformation_output}[{y_label}]"
        y = apply_custom_transformation(y, transformation_output)
        st.write(f"Applied {transformation_output} transformation to {y_label}")

    # Plot original data
    plt.figure(figsize=(10, 6))
    plt.scatter(x, y, color='blue', label='Original Data')

    # Add threshold if given
    if threshold:
        plt.axhline(y=threshold, color='red', linestyle='--', label=f'Threshold = {threshold}')
    
    plt.xlabel(x_label)
    plt.ylabel(y_label)
    plt.title(f'{y_label} vs {x_label} with Threshold')
    plt.legend()
    st.pyplot(plt)
    return x_label, y_label

# Function to plot â vs. a analysis with confidence intervals
def plot_a_hat_vs_a(x, y, model):
    plt.figure(figsize=(10, 6))
    plt.scatter(x, y, color='blue', label='Data')

    # Fitted line and confidence intervals
    predictions = model.predict(add_constant(x))
    plt.plot(x, predictions, color='red', label='Fitted line')

    st.write("Confidence Interval for Predictions:")
    pred_se = np.sqrt(model.mse_resid)
    ci_lower = predictions - 1.96 * pred_se
    ci_upper = predictions + 1.96 * pred_se
    plt.fill_between(x, ci_lower, ci_upper, color='gray', alpha=0.3, label='95% Confidence Interval')

    plt.xlabel(f'Input Parameter (a)  [{x_label}]')
    plt.ylabel(f'Output Parameter (â)  [{y_label}]')
    plt.title('â vs. a Analysis with Confidence Intervals')
    plt.legend()
    st.pyplot(plt)
    
    y_binary_predictions= (predictions >= threshold_input).astype(int)
    st.write(
        f"_______________________________________________________________________________________ <br><br> "
        f"Now performing binary output data calculation using the linear fitted line, followed by logistic POD analysis to determine the POD curve. Below is a summary of the model: <br> ",
        unsafe_allow_html=True
    )
    
    logistic_pod(x.reshape(-1,1), y_binary_predictions.reshape(-1,1))

# Function to plot LS PODanalysis with confidence intervals
def plot_LS_POD(x, y, model):
    plt.figure(figsize=(10, 6))
    plt.scatter(x, y, color='blue', label='Data')

    # Fitted line and confidence intervals
    predictions = model.predict(add_constant(x))
    plt.plot(x, predictions, color='red', label='Fitted line')

    st.write("Confidence Interval for Predictions:")
    pred_se = np.sqrt(model.mse_resid)
    ci_lower = predictions - 1.96 * pred_se
    ci_upper = predictions + 1.96 * pred_se
    plt.fill_between(x, ci_lower, ci_upper, color='gray', alpha=0.3, label='95% Confidence Interval')

    plt.xlabel(f'Input Parameter (a)  [{x_label}]')
    plt.ylabel(f'Output Parameter (â)  [{y_label}]')
    plt.title('LS POD Analysis with Confidence Intervals')
    plt.legend()
    st.pyplot(plt)
    
    y_binary_predictions= (predictions >= threshold_input).astype(int)
    st.write(
        f"_______________________________________________________________________________________ <br><br> "
        f"Now performing binary output data calculation using the linear fitted line, followed by logistic POD analysis to determine the POD curve. Below is a summary of the model: <br> ",
        unsafe_allow_html=True
    )
    
    logistic_pod(x.reshape(-1,1), y_binary_predictions.reshape(-1,1))



# Function to plot logistic POD with confidence intervals
def plot_logistic_pod(x, y_binary, model, scaler):
    plt.figure(figsize=(10, 6))
    plt.scatter(x, y_binary, color='blue', label='Data')

    # Generate x_range within actual x values
    x_range = np.linspace(min(x), max(x), 1000).reshape(-1, 1)

    # Scale x_range for prediction
    x_range_scaled = scaler.transform(x_range)
    prob_predictions = model.predict_proba(x_range_scaled)[:, 1]

    # Plot the POD curve with actual x values
    plt.plot(x_range, prob_predictions, color='red', label='POD curve')

    # Find the exact x value where the prediction is exactly 0.9 using interpolation
    interp_function = interp1d(prob_predictions, x_range.flatten())
    try:
        x_exact = interp_function(0.9)
    except ValueError:
        st.write("No exact x found where prediction is exactly 0.9 within the given range.")
        return

    # Plot the point where POD = 0.9
    plt.scatter([x_exact], [0.9], color='black', zorder=5)
    plt.text(x_exact, 0.9, f'({x_exact:.4f}, 0.9)', fontsize=12, ha='right', va='bottom')

    # Horizontal and vertical lines
    plt.axhline(y=0.9, color='green', linestyle='--', label='POD = 0.9')
    plt.axvline(x=x_exact, color='purple', linestyle='--', label=f'Input Parameter = {x_exact:.4f}')

    plt.xlabel(f'Input Parameter [{x_label}]')
    plt.ylabel(f'Probability of Detection (POD)')
    plt.title('Logistic POD Analysis')
    plt.legend()
    st.pyplot(plt)

# New Function to plot actual output vs input with threshold line
def plot_actual_vs_input(x, y, threshold):
    plt.figure(figsize=(10, 6))
    plt.scatter(x, y, color='blue', label='Actual Output Data')

    # Plotting the threshold line
    plt.axhline(y=threshold, color='red', linestyle='--', label=f'Threshold = {threshold}')

    plt.xlabel(f'Input Parameter (a) [{x_label}]')
    plt.ylabel(f'Output Parameter [{y_label}]')
    plt.title('Actual Output vs Input Data with Threshold Line')
    plt.legend()
    st.pyplot(plt)

# Function to perform and plot â vs. a analysis
def a_hat_vs_a(x, y):
    x0=x
    x = add_constant(x)
    model = OLS(y, x).fit()

    st.write("â vs. a Analysis Summary:")
    st.write(model.summary())

    residuals = y - model.predict(x)
    check_normality_residuals(residuals)
    
    plot_a_hat_vs_a(x0, y, model)

    # Function to perform and plot â vs. a analysis
def LS_POD(x, y):
    x0=x
    x = add_constant(x)
    model = OLS(y, x).fit()

    st.write("LS POD Analysis Summary:")
    st.write(model.summary())

    residuals = y - model.predict(x)
    check_normality_residuals(residuals)
    
    plot_LS_POD(x0, y, model)

# Function to perform and plot Logistic POD analysis with scaled input
def logistic_pod(x, y_binary):
    
    # Scale the input for better logistic regression fitting
    scaler = StandardScaler()
    x_scaled = scaler.fit_transform(x)

    model = LogisticRegression()
    X_train, X_test, y_train, y_test = train_test_split(x_scaled, y_binary, test_size=0.2)
    #model.fit(X_train, y_train)
    model.fit(x_scaled, y_binary)
    y_pred = model.predict(X_test)

    st.write("Logistic Regression Data Analysis:")
    st.write(
        f"P(y=1|X) = 1 / (1 + e^(-(β0 + β1*X1 + β2*X2 + ... + βn*Xn))) where: <br><br> "
        f"P(y=1|X): The probability that the output variable (y) is 1, given the input features (X) <br> "
        f"β0: The intercept term<br> "
        f"β1, β2, ..., βn: The coefficients for the input features (X1, X2, ..., Xn)<br> "
        f"e: The base of the natural logarithm (approximately 2.718).",
        unsafe_allow_html=True
    )

    # st.write("Logistic POD Analysis:")
    # st.write(f"**Coefficients (β1):** {model.coef_}")
    # st.write(f"**Intercept (β0):** {model.intercept_}")

    # st.write("Classification Report:")

    # # Generate the classification report
    # report_dict = classification_report(y_test, y_pred, output_dict=True)
    # report_df = pd.DataFrame(report_dict).transpose()

    # # Display the DataFrame in Streamlit
    # st.dataframe(report_df)

    plot_logistic_pod(x, y_binary, model,scaler)

# Function to perform and plot Hit/Miss Analysis
def hit_miss(x, y_binary):
    # Scale the input for better logistic regression fitting
    scaler = StandardScaler()
    x_scaled = scaler.fit_transform(x)

    model = LogisticRegression().fit(x_scaled, y_binary)
    st.write("Hit/miss POD uses Logistic Regression Data Analysis:")
    st.write(
        f"P(y=1|X) = 1 / (1 + e^(-(β0 + β1*X1 + β2*X2 + ... + βn*Xn))) where: <br><br> "
        f"P(y=1|X): The probability that the output variable (y) is 1, given the input features (X) <br> "
        f"β0: The intercept term<br> "
        f"β1, β2, ..., βn: The coefficients for the input features (X1, X2, ..., Xn)<br> "
        f"e: The base of the natural logarithm (approximately 2.718).",
        unsafe_allow_html=True
    )
    st.write("Hit/Miss Analysis:")
    st.write(f"Coefficients: {model.coef_}")
    st.write(f"Intercept: {model.intercept_}")

    plot_logistic_pod(x, y_binary, model,scaler)

# Function to perform and report 29/29 Analysis
def perform_29_29(x, y_binary,check_2929):
    detection_rate = y_binary.mean()
    if check_2929 == 'Yes':
        st.write("29/29 Analysis:")
        if detection_rate == 1.0:
            st.write("Detection method meets 29/29 criteria.")
        else:
            st.write("Detection method does not meet 29/29 criteria.")
    elif check_2929 == 'No':
        st.write("This data/test does not satisfy the requirement for POD 29/29")
    
    k =sum(y_binary)
    n=len(y_binary)
    alpha=0.95
    lower_bound = binom.ppf(alpha / 2, n, k / n) / n
    upper_bound = binom.ppf(1 - alpha / 2, n, k / n) / n
    st.write(f"Clopper-Pearson Interval: [{lower_bound:.4f}, {upper_bound:.4f}]")



# Function to check normality of input data before analysis
def check_normality_data(x, y):
    st.write("Checking Normality of Data:")

    fig, ax = plt.subplots(2, 2, figsize=(12, 12))

    # Histogram of input data
    ax[0, 0].hist(x, bins=30, color='lightgreen')
    ax[0, 0].set_title(f'Histogram of Input Data [{x_label}]')

    # Residual plot of output as a function of input
    x_with_const = add_constant(x)
    model = OLS(y, x_with_const).fit()
    residuals = y - model.predict(x_with_const)
    ax[0, 1].scatter(x, residuals, color='blue')
    ax[0, 1].axhline(y=0, color='red', linestyle='--')
    ax[0, 1].set_title(f'Residuals vs Input Data [{x_label}]')
    ax[0, 1].set_xlabel(f'Input Data [{x_label}]')
    ax[0, 1].set_ylabel('Residuals')

    # Probability Plot of Standardized Residuals with Confidence Intervals
    standardized_residuals = (residuals - np.mean(residuals)) / np.std(residuals)
    # hist plot of standardized residuals
    ax[1, 0].hist(standardized_residuals,color='lightgreen')
    ax[1, 0].set_title('Histogram of of standardized Residuals')

    (osm, osr), (slope, intercept, r) = stats.probplot(standardized_residuals, dist="norm")
    
    ax[1, 1].plot(osm, osr, 'o', label='Standardized Residuals')
    ax[1, 1].plot(osm, slope * osm + intercept, 'r-', label='Fit Line')
    
    # Confidence Intervals
    ci = 1.96 / np.sqrt(len(standardized_residuals)) * np.std(standardized_residuals)
    ax[1, 1].fill_between(osm, slope * osm + intercept - ci, slope * osm + intercept + ci, color='gray', alpha=0.2, label='95% Confidence Interval')
    
    ax[1, 1].set_title('Probability Plot of Standardized Residuals')
    ax[1, 1].set_xlabel('Theoretical Quantiles')
    ax[1, 1].set_ylabel('Standardized Residuals')
    ax[1, 1].legend()
    
    st.pyplot(fig)

    # Shapiro-Wilk test for normality
    shapiro_test = stats.shapiro(residuals)
    st.write(f"Shapiro-Wilk Test: W={shapiro_test[0]:.4f}, p-value={shapiro_test[1]:.4f}")

    if shapiro_test[1] < 0.05:
        st.write("The residuals do not follow a normal distribution (reject H0).")
    else:
        st.write("The residuals follow a normal distribution (fail to reject H0).")

# Function to save content to a Word document
def save_to_word(summary_text, plots):
    doc = Document()

    # Add summary text
    doc.add_heading('POD Data Analysis Report', 0)
    doc.add_paragraph(summary_text)

    # Add plots
    for plt in plots:
        image_stream = BytesIO()
        plt.savefig(image_stream, format='png')
        image_stream.seek(0)
        doc.add_picture(image_stream, width=Inches(6))

    # Save the document
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer




# _______________________________________________________________________________________________

# POD Data Analysis Streamlit App
st.title("POD Data Analysis App")

# Sidebar layout with two columns for data analysis and method selection
st.sidebar.title("Data Analysis Configuration")
uploaded_file = st.file_uploader("Upload your data file (Excel)", type=["xlsx"])

if uploaded_file:
    try:
        data = pd.read_excel(uploaded_file, engine='openpyxl')
        st.write("File successfully read.")
    except Exception as e:
        st.error(f"Failed to read the file: {e}")
        data = None

    if data is not None and not data.empty:
        st.write("Data Preview:")
        st.dataframe(data.head())

        valid_columns = data.columns.tolist()

        if len(valid_columns) < 2:
            st.error("The file doesn't have enough columns for analysis. Please check the file format.")
        else:
            input_param = st.selectbox("Select Input Parameter:", valid_columns)
            output_param = st.selectbox("Select Output Parameter:", valid_columns)
            

            if input_param and output_param:
                data[input_param] = pd.to_numeric(data[input_param], errors='coerce')
                data[output_param] = pd.to_numeric(data[output_param], errors='coerce')
                

                cleaned_data = data.dropna(subset=[input_param, output_param])
                

                if cleaned_data.empty :
                    st.error("After cleaning, the dataset is empty. Please check your input parameters and data.")
                else:
                    # Sidebar inputs for analysis configuration
                    threshold_input = st.sidebar.number_input("Enter rejection threshold", value=0.5)
                    method = st.sidebar.selectbox("Select Analysis Method", ["LS POD","â vs a", "Logistic POD", "Hit/Miss", "29/29"])

                    # Full data analysis based on input
                    x = cleaned_data[input_param].values
                    y = cleaned_data[output_param].values
                    output_set = set(y)

                    if (output_set.issubset({0, 1})):
                        check_binay_y = st.selectbox(f"looks like the output data are binary?", ["Yes", "No"])
                    else:
                        check_binay_y='No'

                    # Always show full data analysis plot
                    [x_label,y_label]=full_data_analysis(x, y, input_param, output_param, threshold_input)

                    # Checking normality of input data
                    if check_binay_y=='No':
                        check_normality_data(cleaned_data[input_param],cleaned_data[output_param])

                    if method== "29/29":
                        check_2929 = st.selectbox(f"Please confirm all input parameters is identical and all output data is correspond to same flaw size", ["","Yes","No"])

                    if st.sidebar.button("Perform Analysis"):
                        if check_binay_y=='No':
                            y_binary = (y >= threshold_input).astype(int)
                        else:
                            y_binary = (y>= 0.5).astype(int)
                        
                        
                

                        # Perform the selected analysis
                        if method == "â vs a":
                            a_hat_vs_a(x, y)
                        elif method == "Logistic POD":
                            logistic_pod(x.reshape(-1, 1), y_binary)
                        elif method == "Hit/Miss":
                            hit_miss(x.reshape(-1, 1), y_binary)
                        elif method == "29/29":
                            perform_29_29(x, y_binary,check_2929)
                        elif method == "LS POD":
                            LS_POD(x, y)
                    
                        # summary_text = ""

                        # if check_binay_y == 'No':
                        #     y_binary = (y >= threshold_input).astype(int)
                        # else:
                        #     y_binary = (y >= 0.5).astype(int)

                        # # Perform the selected analysis
                        # plots = []
                        # if method == "â vs a":
                        #     a_hat_vs_a(x, y)  # This function will also generate the plots
                        #     summary_text += "â vs a Analysis Summary.\n"
                        #     plots.append(plt)
                        # elif method == "Logistic POD":
                        #     logistic_pod(x.reshape(-1, 1), y_binary)
                        #     summary_text += "Logistic POD Analysis Summary.\n"
                        #     plots.append(plt)
                        # elif method == "Hit/Miss":
                        #     hit_miss(x.reshape(-1, 1), y_binary)
                        #     summary_text += "Hit/Miss Analysis Summary.\n"
                        #     plots.append(plt)
                        # elif method == "29/29":
                        #     perform_29_29(x, y_binary, check_2929)
                        #     summary_text += "29/29 Analysis Summary.\n"
                        #     plots.append(plt)
                        # elif method == "LS POD":
                        #     LS_POD(x, y)
                        #     summary_text += "LS POD Analysis Summary.\n"
                        #     plots.append(plt)

                        # # Save to Word and provide download link
                        # buffer = save_to_word(summary_text, plots)

                        # st.download_button(
                        #     label="Download Report",
                        #     data=buffer,
                        #     file_name="POD_Analysis_Report.docx",
                        #     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        # )
                            
else:
    st.write("Please upload a data file to get started.")
